"""
LLA (Leave & License Agreement) Word document generator.
Uses python-docx to create .docx files server-side.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
from datetime import datetime


def _ordinal(day):
    if 4 <= day <= 20:
        return f"{day}th"
    return f"{day}{['th','st','nd','rd'][day % 10] if day % 10 < 4 else 'th'}"


def _format_date(date_str):
    if not date_str:
        return ""
    try:
        d = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        months = ["January","February","March","April","May","June",
                   "July","August","September","October","November","December"]
        return f"{d.day} {months[d.month-1]}, {d.year}"
    except Exception:
        return date_str


def _month_year(date_str):
    if not date_str:
        return ""
    try:
        d = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        months = ["January","February","March","April","May","June",
                   "July","August","September","October","November","December"]
        return f"{months[d.month-1]} {d.year}"
    except Exception:
        return date_str


def _calc_end_date(start_str, months=11):
    if not start_str:
        return ""
    try:
        d = datetime.fromisoformat(start_str.replace('Z', '+00:00'))
        month = d.month + months
        year = d.year + (month - 1) // 12
        month = ((month - 1) % 12) + 1
        day = min(d.day, 28)
        end = datetime(year, month, day) 
        return end.strftime('%Y-%m-%d')
    except Exception:
        return ""


def _fmt_inr(amount):
    if not amount:
        return "0"
    n = int(amount)
    s = str(n)
    if len(s) > 3:
        last3 = s[-3:]
        rest = s[:-3]
        groups = []
        while rest:
            groups.insert(0, rest[-2:])
            rest = rest[:-2]
        return ','.join(groups) + ',' + last3
    return s


def generate_lla_docx(company: dict) -> bytes:
    """Generate a Leave & License Agreement .docx from company data."""
    
    # Build agreement data
    end_date = company.get('end_date') or _calc_end_date(company.get('start_date'), company.get('lock_in_months', 11))
    security = company.get('security_deposit') or (company.get('total_seats', 1) * company.get('rate_per_seat', 0))
    
    data = {
        'day': _ordinal(datetime.fromisoformat(company.get('start_date', '2026-01-01').replace('Z', '+00:00')).day) if company.get('start_date') else '1st',
        'month_year': _month_year(company.get('start_date')),
        'company_name': company.get('company_name', ''),
        'pan': company.get('signatory_pan') or company.get('company_pan', ''),
        'gstin': company.get('company_gstin', 'N/A'),
        'address': company.get('company_address', ''),
        'signatory': company.get('signatory_name', ''),
        'designation': company.get('signatory_designation', 'Authorized Signatory'),
        'space': company.get('space_description') or company.get('plan_name', 'Workspace'),
        'seats': company.get('total_seats', 1),
        'start_date': _format_date(company.get('start_date')),
        'end_date': _format_date(end_date),
        'lock_in': company.get('lock_in_months', 11),
        'fee': company.get('rate_per_seat', 0),
        'security': security,
        'setup': company.get('setup_charges', 'Not applicable'),
    }

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Helper functions
    def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = space_after
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return p

    def add_heading_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return p

    def add_numbered(num, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"{num} ")
        run_num.bold = True
        run_num.font.name = 'Calibri'
        run_num.font.size = Pt(11)
        run_text = p.add_run(text)
        run_text.font.name = 'Calibri'
        run_text.font.size = Pt(11)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.text = ""
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return p

    def add_mixed_para(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = space_after
        for text, bold in parts:
            run = p.add_run(text)
            run.bold = bold
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        return p

    # === TITLE ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run("LEAVE & LICENSE AGREEMENT")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(14)

    # === OPENING ===
    add_mixed_para([
        ('THIS LEAVE AND LICENSE AGREEMENT ("Agreement") is executed on this ', False),
        (data['day'], True),
        (' day of ', False),
        (data['month_year'], True),
        (', at Faridabad, Haryana,', False),
    ])

    # BETWEEN
    add_para("BETWEEN", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Licensor
    add_mixed_para([
        ('Thryve Coworking', True),
        (', (PAN AAYFT8213A & GSTIN 06AAYFT8213A1Z2) a business operating from: Plot No. 3, First Floor, Near Ajronda Metro Station, 18/1, Mathura Road, Faridabad, Haryana 121007, acting through its authorized representative, hereinafter referred to as the "', False),
        ('Licensor', True),
        ('", (which expression shall, unless repugnant to the context, include its successors, assigns, administrators, and representatives)', False),
    ])

    add_para("AND", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # Licensee
    add_mixed_para([
        (f"M/s {data['company_name']}", True),
        (f", (PAN {data['pan']} & GSTIN {data['gstin']}) having its Regd. office at: {data['address']}, through its authorized signatory ", False),
        (data['signatory'], True),
        (', hereinafter referred to as the "', False),
        ('Licensee', True),
        ('", (which expression shall, unless repugnant to the context, include its successors, permitted assigns, employees, and representatives)', False),
    ])

    add_para('The Licensor and the Licensee are collectively referred to as the "Parties" and individually as a "Party."')

    # WHEREAS
    add_heading_text("WHEREAS")
    add_numbered("1.", 'The Licensor is operating a coworking facility known as "Thryve Coworking" at Plot No. 3, First Floor, Near Ajronda Metro Station, 18/1, Mathura Road, Faridabad 121007 ("hereinafter referred to as Premises").')
    add_numbered("2.", "The Licensor is in lawful possession of the Premises and entitled to grant a license for the use of certain designated workspace(s), cabins, meeting rooms, and/or coworking desks.")
    add_numbered("3.", "The Licensee has requested use of the Licensor's coworking facilities and services, and the Licensor has agreed to grant the Licensee a limited, non-exclusive, revocable license with rights to enter and strictly for business use on the terms and conditions set forth below:")
    
    add_para("NOW, THEREFORE, THE PARTIES AGREE AS FOLLOWS:", bold=True)

    # Section 1
    add_heading_text("1. LICENSED AREA & SERVICES")
    add_numbered("1.1", "The Licensor grants to the Licensee a non-exclusive, non-transferable, revocable license for right to enter and use the following workspace(s) at the Premises:")
    add_bullet(data['space'])
    add_bullet(f"Number of Seats: {data['seats']}")
    add_bullet("Meeting Room Access: As per usage/plan")
    add_bullet("Common Areas: Reception, Lounge, Pantry, Washrooms, Hallways, and such other areas as designated by the Licensor.")
    add_numbered("1.2", "The License is limited to right to enter and use only, and no tenancy, lease, or rights of possession are created or transferred.")
    add_numbered("1.3", "The following services are included (as applicable to the chosen plan):")
    for svc in ["High-speed Internet", "Electricity & Air Conditioning during business hours", "Housekeeping of common spaces", "Power Backup", "Front Desk Assistance", "Access to Pantry Facilities", "Access to Conference Rooms (subject to availability and booking rules)", "CCTV Surveillance (excluding private cabins)"]:
        add_bullet(svc)
    add_numbered("1.4", "The Licensor may modify or enhance services at its discretion.")

    # Section 2
    add_heading_text("2. TERM")
    add_mixed_para([
        ("2.1 ", True),
        ("The term of this Agreement shall be for 11 months, commencing from ", False),
        (data['start_date'], True),
        (' ("Start Date") until ', False),
        (data['end_date'], True),
        (' ("End Date"), unless terminated earlier. There will be a lock-in period of ', False),
        (f"{data['lock_in']} months", True),
        (" before which the Licensee will not terminate this Leave & License Deed.", False),
    ])
    add_numbered("2.2", "Renewal shall be subject to discretion of the Licensor and a new Leave & License Deed will be executed.")

    # Section 3
    add_heading_text("3. LICENSE FEES & PAYMENTS")
    add_numbered("3.1", "The Licensee agrees to pay:")
    add_bullet(f"Monthly License Fee per seat/member: Rs. {_fmt_inr(data['fee'])}/-")
    add_bullet(f"Security Deposit (interest-free): Rs. {_fmt_inr(data['security'])}/-")
    setup = f"Rs. {_fmt_inr(data['setup'])}/-" if isinstance(data['setup'], (int, float)) else str(data['setup'])
    add_bullet(f"Setup Charges (if any): {setup}")
    add_bullet("GST/other applicable taxes/extra usage of facilities will be paid extra.")
    add_numbered("3.2", "The License Fee shall be payable in advance on or before the 5th day of each English calendar month.")
    add_numbered("3.3", "Delay in payment beyond the due date attracts a penalty of Rs. 500/- per day apart from the outstanding dues until the full outstanding payment is made.")
    add_numbered("3.4", "Non-payment for two consecutive months constitutes a material breach, allowing the Licensor to terminate the Agreement without further notice.")

    # Sections 4-15
    for title_text, items in [
        ("4. SECURITY DEPOSIT", [
            ("4.1", "The Security Deposit shall be refunded within 30 working days from the date of vacating the Premises, after adjusting:"),
            None, "Pending dues", "Damages (if any)", "Loss of property or assets", "Outstanding penalties",
            ("4.2", "The Security Deposit shall not be used by the Licensee towards monthly fees."),
        ]),
        ("5. USE OF PREMISES", [
            ("5.1", "The Premises shall be used only for lawful business activities. Any illegal, immoral, or hazardous activities are strictly prohibited."),
            ("5.2", "The Licensee shall not:"),
            None, "Install permanent fixtures", "Alter or damage the Premises", "Store hazardous materials", "Conduct activities causing nuisance, noise, or disturbance", "Allow unauthorized persons to occupy the workspace",
            ("5.3", "The Licensee must maintain cleanliness and hygiene in its designated area."),
            ("5.4", "Sub-licensing, assignment, or sharing of the workspace is strictly prohibited."),
        ]),
        ("6. MEETING ROOM & FACILITY RULES", [
            ("6.1", "Meeting room usage is subject to prior booking and availability."),
            ("6.2", "Overstay beyond the reserved slot may attract overtime charges."),
            ("6.3", "Common area usage shall be respectful and non-disruptive. Any events or visitors should be notified in advance and are subject to approval & charges at the discretion of the Licensor."),
            ("6.4", "Pantry use is for light refreshments only; cooking/heating heavy food items is prohibited."),
        ]),
        ("7. INTERNET, IT & ELECTRICAL USAGE", [
            ("7.1", "The Licensee shall use the Licensor's internet responsibly and not engage in illegal downloads, hacking, or bandwidth abuse."),
            ("7.2", "High electrical load equipment (servers, printers, etc.) requires prior approval."),
            ("7.3", "The Licensor is not liable for internet downtime caused by third-party providers."),
        ]),
        ("8. ACCESS & SECURITY", [
            ("8.1", "Normal access hours: Monday to Saturday 9:00 AM to 9:00 PM. The Premises will remain closed on all Sundays and Govt. declared public holidays (as notified from time to time) unless otherwise permitted."),
            ("8.2", "Licensee shall not tamper with CCTV, access control systems, or security devices."),
            ("8.3", "Visitors must follow the Licensor's entry protocols."),
        ]),
        ("9. DAMAGE & LOSS", [
            ("9.1", "The Licensee shall be responsible for any damage caused by its employees, agents, or guests."),
            ("9.2", "The Licensor is not responsible for theft, loss, or damage to personal belongings, laptops, or equipment brought by the Licensee."),
        ]),
        ("10. TERMINATION", [
            ("10.1", "Either Party may terminate this Agreement by giving 30 days prior written notice."),
            ("10.2", "The Licensor may terminate this leave and license agreement immediately if there is:"),
            None, "Non-payment of license fee for two months", "Any illegal activity, misconduct or nuisance by the Licensee or its employees", "Any damage caused to property & person by the Licensee or its employees", "Any Breach of Agreement terms by the Licensee or its employees",
            ("10.3", "On termination, the Licensee shall:"),
            None, "Vacate the workspace peacefully", "Remove all personal belongings", "Return access cards/keys", "Clear all dues",
        ]),
        ("11. NO TENANCY / NO LEASE", [
            ("11.1", "This Agreement creates no tenancy, lease, occupancy rights, or interest in the Premises."),
            ("11.2", "The Licensee acknowledges that:"),
            None, "The Licensor retains full control over the Premises", "The Licensee's right is purely permissive", "This Agreement does not fall under the Rent Control Act",
        ]),
        ("12. INDEMNITY", [
            ("12.1", "The Licensee agrees to indemnify and hold harmless the Licensor from all:"),
            None, "Legal claims", "Damages", "Liabilities", "Losses arising out of the Licensee's use of the Premises.",
            ("12.2", "The Licensee will be solely responsible for compliance with all applicable laws, payments of all statutory levies, dues, or liabilities including Goods & Service Tax (GST) arising from the use of the licensed premises. The Licensor will not be held responsible in any manner whatsoever for any such liabilities, defaults or non-compliances on the part of the Licensee."),
        ]),
        ("13. LIMITATION OF LIABILITY", [
            ("13.1", "The Licensor shall not be responsible for:"),
            None, "Business losses", "Data loss", "Interruption of services due to external factors", "Acts of God, electrical failures, or internet breakdowns",
            ("13.2", "The maximum liability of the Licensor shall not exceed the monthly license fee."),
        ]),
        ("14. GOVERNING LAW & JURISDICTION", [
            ("14.1", "This Agreement shall be governed by the laws of India."),
            ("14.2", "Courts in Faridabad, Haryana shall have exclusive jurisdiction."),
        ]),
        ("15. MISCELLANEOUS", [
            ("15.1", "Any amendments must be in writing and signed by both Parties."),
            ("15.2", "Notices shall be served via email and registered post."),
            ("15.3", "If any clause is held invalid, the remaining clauses shall remain enforceable."),
        ]),
    ]:
        add_heading_text(title_text)
        in_bullets = False
        for item in items:
            if item is None:
                in_bullets = True
                continue
            if isinstance(item, tuple):
                in_bullets = False
                add_numbered(item[0], item[1])
            elif in_bullets:
                add_bullet(item)

    # Section 16 - Signatures
    add_heading_text("16. SIGNATURES")
    add_para("IN WITNESS WHEREOF, the Parties hereto have executed this Agreement on the date and year first written above.")

    doc.add_paragraph()
    add_para("FOR THE LICENSOR", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Thryve Coworking", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Name: Amit Mehta", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Designation: Marketing Head", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Signature: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(f"Date: {data['start_date']}", align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    add_para("FOR THE LICENSEE", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(f"Name of Company: {data['company_name']}", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(f"Authorized Signatory: {data['signatory']}", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(f"Designation: {data['designation']}", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Signature: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(f"Date: {data['start_date']}", align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    add_para("WITNESSES:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    add_para("1. Name: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("   Address: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("   Signature: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    add_para("2. Name: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("   Address: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("   Signature: ________________", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Write to bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
